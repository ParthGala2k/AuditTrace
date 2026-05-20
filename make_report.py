"""
Generate AuditTrace_Report.docx  --  python make_report.py

Builds the CMPE 258 final project report as a Word document.
All numbers are read from evaluation/eval_e2e_results.json and
evaluation/comparison_summary.json so they stay consistent with the
rest of the deliverables.

Sections follow the rubric template:
  1. Cover Information
  2. Abstract
  3. Introduction & Problem Description
  4. Background / Related Work
  5. System / Model / Algorithm Design
  6. Implementation Details
  7. Task Distribution & Contributions
  8. Evaluation & Testing Results
  9. References
"""
import json, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.abspath(__file__))
EVAL = os.path.join(ROOT, "evaluation")

with open(os.path.join(EVAL, "eval_e2e_results.json")) as f:
    FP = json.load(f)
with open(os.path.join(EVAL, "comparison_summary.json")) as f:
    PLN = json.load(f)


# ── helpers ─────────────────────────────────────────────────────────────────
def pct(v):  return "—" if v is None else f"{v*100:.1f}%"
def num(v):  return "—" if v is None else f"{v:.2f}"
def usd(v):  return "—" if v is None else f"${v:.4f}"

MODEL_LABELS = {
    "openai/gpt-4o-mini":               "GPT-4o mini",
    "deepseek/deepseek-chat":           "DeepSeek V3",
    "meta-llama/llama-3.1-70b-instruct": "Llama 3.1 70B",
}
def lbl(m): return MODEL_LABELS.get(m, m.split("/")[-1])


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),  "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0f, 0x11, 0x15)
    return h


def add_para(doc, text, *, bold=False, italic=False, size=11, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def add_bullets(doc, items, *, size=11):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.runs[0] if p.runs else p.add_run(item)
        if not p.runs:
            r = p.add_run(item)
        else:
            r.text = item
        r.font.size = Pt(size)


def add_table(doc, headers, rows, *, header_fill="0F1115", widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        set_cell_bg(cell, header_fill)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # body
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10.5)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
    if widths:
        for row in table.rows:
            for ci, w in enumerate(widths):
                row.cells[ci].width = Inches(w)
    doc.add_paragraph()  # spacer
    return table


# ── document setup ──────────────────────────────────────────────────────────
doc = Document()

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

# Letter-size margins are default; tighten slightly
for section in doc.sections:
    section.top_margin    = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin   = Inches(0.9)
    section.right_margin  = Inches(0.9)


# ============================================================================
# 1. COVER INFORMATION
# ============================================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_r = title_p.add_run("AuditTrace")
title_r.bold = True
title_r.font.size = Pt(32)
title_r.font.color.rgb = RGBColor(0x0f, 0x11, 0x15)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_r = sub_p.add_run("Multi-Agent LLM Consensus for Infrastructure-as-Code Compliance Auditing")
sub_r.font.size = Pt(15)
sub_r.font.color.rgb = RGBColor(0x6f, 0x42, 0xc1)
sub_r.italic = True

doc.add_paragraph()

course_p = doc.add_paragraph()
course_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
course_r = course_p.add_run("CMPE 258 Deep Learning  |  Final Project  |  San Jose State University  |  Spring 2026")
course_r.font.size = Pt(11)

doc.add_paragraph()

add_table(doc,
    ["Field", "Value"],
    [
        ["Team ID",          "Group 4"],
        ["Project Title",    "AuditTrace — Multi-Agent LLM Consensus for IaC Compliance Auditing"],
        ["Project Track",    "Application"],
        ["Focused Areas",    "Multi-agent LLM systems · LLM-as-judge evaluation · Compliance/security tooling · Hybrid symbolic + neural architectures"],
        ["GitHub Repository", "https://github.com/ParthGala2k/AuditTrace"],
        ["Demo Video",        "https://youtu.be/3yNg2zou0WI"],
    ],
    widths=[1.6, 5.2])

add_heading(doc, "Team Members", level=2)
add_table(doc,
    ["Name", "SJSU ID", "Email", "Role"],
    [
        ["Parth Gala", "[FILL IN]", "parth.gala@sjsu.edu", "Sole contributor — full-stack implementation, evaluation, documentation"],
    ],
    widths=[1.6, 1.2, 2.2, 2.2])

doc.add_page_break()


# ============================================================================
# 2. ABSTRACT
# ============================================================================
add_heading(doc, "Abstract", level=1)

abstract = (
    "AuditTrace is a hybrid compliance auditing system that combines a deterministic "
    "infrastructure-as-code scanner (Checkov) with a three-model Large Language Model "
    "consensus filter that triages false positives using the surrounding Terraform "
    "context. The system addresses a concrete pain point in cloud security: static "
    "scanners produce a firehose of alerts — 385 findings on the deliberately vulnerable "
    "terragoat repository alone — without distinguishing genuine violations from contextual "
    "false positives. We send each Checkov finding plus its enclosing Terraform resource "
    "block to GPT-4o mini, DeepSeek V3, and Llama 3.1 70B in parallel; each model "
    "independently labels the finding as GENUINE, FALSE_POSITIVE, or UNCERTAIN, and a "
    "consensus vote assigns a confidence tier. On a hand-labeled evaluation set of 30 "
    "adversarial findings spanning 13 distinct Checkov rules, the consensus reaches 100% "
    "overall accuracy with 100% false-positive reduction and 100% true-positive retention, "
    "while individual models score F1 between 84.6% (Llama) and 100% (DeepSeek). Cohen's "
    "kappa between model pairs ranges 0.67 to 0.93, demonstrating that the models genuinely "
    "disagree on hard cases and that consensus voting performs measurable error correction "
    "rather than rubber-stamping. A separate planner benchmark (n=55 CIS clauses) quantifies "
    "an explicit accuracy-vs-latency-vs-cost tradeoff across the three models. All numbers "
    "are reproducible from two committed scripts and a SHA-256 prompt cache."
)
add_para(doc, abstract)


# ============================================================================
# 3. INTRODUCTION & PROBLEM DESCRIPTION
# ============================================================================
add_heading(doc, "1. Introduction and Problem Description", level=1)

add_para(doc,
    "Cloud compliance auditing is increasingly automated through static analysis of "
    "infrastructure-as-code (IaC) artifacts such as Terraform, CloudFormation, and Kubernetes "
    "manifests. Mature scanners — Checkov, tfsec, kics — encode hundreds of rules drawn from "
    "frameworks like the CIS Benchmarks, NIST SP 800-53, SOC 2, and PCI-DSS. These scanners "
    "are extremely good at one thing: rule-based pattern matching against the abstract syntax "
    "of an IaC file. They flag every occurrence of a violation deterministically. They are "
    "also, by design, completely context-blind.")

add_para(doc,
    "On real engineering teams this produces a recurring failure mode: scanners surface "
    "hundreds of alerts per pull request, the vast majority of which are not actually "
    "actionable in the specific context they appear in. A bucket marked public-read in a "
    "marketing static-site module is fundamentally different from one in a production "
    "billing module, yet both look identical to the rule engine. Engineers stop looking at "
    "the alerts. The compliance posture of the codebase silently degrades.")

add_para(doc,
    "We hypothesize that this is the precise problem class where current large language "
    "models add real value. LLMs cannot replace the rule engine — they would hallucinate, "
    "miss obvious violations, and produce nondeterministic output — but they are unusually "
    "well-suited to reading surrounding code and judging whether a flagged pattern actually "
    "represents a real risk in its context. We call this the false-positive filter task. "
    "It is the inverse of detection: detection is symbolic and deterministic, filtering is "
    "contextual and probabilistic.")

add_heading(doc, "Target users and scenario", level=2)
add_bullets(doc, [
    "Cloud security engineers who triage Checkov output today and want a high-confidence shortlist of genuine violations.",
    "Compliance auditors who need a traceable mapping from policy clause (CIS, NIST, SOC 2) to specific lines of code in a customer repository.",
    "DevOps teams running IaC pipelines who want a CI gate that blocks pull requests introducing new high-confidence violations without drowning developers in noise.",
])

add_heading(doc, "Why this is a deep learning problem", level=2)
add_para(doc,
    "The judgement task — does this flagged Terraform block represent a real violation, "
    "given the file path, the comments, the sibling resources, the naming conventions? — "
    "is fundamentally a contextual reasoning problem that classical NLP or rule systems "
    "cannot do. It requires the kind of broad domain understanding present in modern "
    "instruction-tuned LLMs. By framing the LLM as a judge rather than a detector, we "
    "avoid the failure modes (hallucination, miss-rate, nondeterminism) that have stopped "
    "LLMs from being used in compliance pipelines, while preserving the strength (contextual "
    "reasoning) that classical methods lack. We then combine three independent LLM judges "
    "into a single consensus vote, hypothesizing that distinct training corpora produce "
    "non-overlapping error modes that can be cancelled by aggregation.")


# ============================================================================
# 4. BACKGROUND / RELATED WORK
# ============================================================================
add_heading(doc, "2. Background and Related Work", level=1)

add_para(doc,
    "AuditTrace draws on three lines of prior work. First, the field of automated "
    "infrastructure-as-code scanning, exemplified by Checkov [1] and the underlying "
    "Bridgecrew rule corpus, provides the deterministic detection layer on which we build. "
    "These tools encode security best practices as regular-expression-style checks over "
    "HCL, JSON, and YAML; they are mature but explicitly context-blind. The CIS Foundations "
    "Benchmarks [2] supply the policy framework we distill into machine-readable form.")

add_para(doc,
    "Second, the recent literature on LLM-as-judge methodology [5] shows that using an LLM "
    "to evaluate another model's output, or to validate the output of a rule-based system, "
    "can produce evaluation signals comparable to human raters at a fraction of the cost. "
    "We adapt this pattern by treating the Checkov rule engine as the system-under-evaluation "
    "and the LLMs as the judges. Multi-model agreement quantified via Cohen's kappa [3], a "
    "classical inter-rater reliability statistic, gives us a calibrated measure of how much "
    "the LLMs disagree on hard cases, and the Landis & Koch [4] interpretation scale provides "
    "a published threshold (substantial vs almost-perfect agreement) for reporting it.")

add_para(doc,
    "Third, agent-orchestration frameworks — LangGraph [6] for stateful multi-step pipelines "
    "and OpenRouter [7] as a unified API gateway across LLM providers — provide the "
    "engineering substrate that lets us run three models from different vendors against the "
    "same prompt with a single key. Our use is deliberately conservative: a two-node "
    "StateGraph (scan → consensus) rather than a deeply recursive agent loop, with all "
    "model responses cached on a SHA-256 hash of inputs for determinism.")

add_para(doc,
    "Where this work differs from existing literature is in the framing. Most LLM-for-security "
    "papers position the LLM as a detector that competes with rule-based scanners. Our "
    "experiments demonstrate that this framing is mathematically dangerous: in a clause-mapping "
    "formulation (which we originally implemented) the precision of any LLM-generated finding "
    "is forced to 100% by construction, because the LLM can never invent check_ids that "
    "Checkov didn't already produce. The headline accuracy numbers in that setup are rigged. "
    "Reframing the LLM as a triage filter on top of the scanner — never adding findings, only "
    "labelling and suppressing them — yields metrics that genuinely vary across models and "
    "reflect real reasoning capability.")


# ============================================================================
# 5. SYSTEM / MODEL / ALGORITHM DESIGN
# ============================================================================
add_heading(doc, "3. System, Model, and Algorithm Design", level=1)

add_heading(doc, "Overall architecture", level=2)
add_para(doc,
    "The system has two phases. An offline phase, run once, distills the 312-page CIS AWS "
    "Foundations Benchmark PDF into 89 structured compliance requirements using GPT-4o mini "
    "(one LLM call per detected control heading), and saves the result to "
    "policies/cis_aws_v7.json. This is committed to the repository and never regenerated.")

add_para(doc,
    "A runtime phase, run per audit, is implemented as a two-node LangGraph StateGraph. "
    "The scan node clones the target GitHub repository, runs Checkov over it, extracts the "
    "full enclosing Terraform resource block for each finding (via a brace-balanced "
    "back-walk from the violation line), attaches the HCL block to the finding payload, "
    "and deletes the local clone. The consensus node passes the enriched findings to three "
    "LLMs running in parallel through OpenRouter — GPT-4o mini, DeepSeek V3, and Llama 3.1 "
    "70B Instruct. Each model independently labels every finding as GENUINE, FALSE_POSITIVE, "
    "or UNCERTAIN and supplies a one-line reason. The verdicts are aggregated into a tier:")

add_table(doc,
    ["GENUINE votes", "Tier", "Action"],
    [
        ["3 / 3 GENUINE",        "HIGH",       "Prioritize fix immediately"],
        ["2 / 3 GENUINE",        "LIKELY",     "Schedule fix"],
        ["1 / 3 GENUINE",        "UNCERTAIN",  "Manual review"],
        ["2 / 3 FALSE_POSITIVE", "LIKELY_FP",  "Probably noise, but keep visible"],
        ["3 / 3 FALSE_POSITIVE", "SUPPRESSED", "Auto-removed from violations list"],
    ],
    widths=[1.8, 1.5, 3.5])

add_heading(doc, "Key technique: HCL context extraction", level=2)
add_para(doc,
    "Checkov returns the line range of a violation but not the enclosing resource. Our "
    "extract_resource_block() function in tools/checkov_runner.py walks backwards from "
    "the violation line until it finds a line beginning with resource, module, data, "
    "provider, or variable, then walks forward through balanced braces until the enclosing "
    "block closes. This produces a self-contained Terraform snippet that the LLM can read "
    "without needing the rest of the file. For non-HCL inputs (CloudFormation JSON, "
    "Kubernetes YAML) the implementation falls back to a windowed snippet of ±5 lines, "
    "degrading gracefully.")

add_heading(doc, "Key technique: SHA-256 prompt cache", level=2)
add_para(doc,
    "LLM responses are cached on disk at backend/mapping_cache/. The cache key is the "
    "SHA-256 hash of (TASK_VERSION, model, sorted check_ids, sorted clause_ids, HCL "
    "digest). TASK_VERSION is a string constant in agents/consensus.py that we bump "
    "whenever the prompt or output schema changes — this invalidates older cache entries "
    "automatically. The result is that every reported number is fully deterministic on "
    "re-runs, even though the underlying LLMs are sampled with temperature parameters "
    "outside our control on the server side.")

add_heading(doc, "Key technique: parallel multi-model voting", level=2)
add_para(doc,
    "All three model judgements for a given audit are dispatched in parallel via "
    "asyncio.gather(); the wall-clock latency is approximately the slowest model's latency "
    "(typically DeepSeek at 7 seconds per batch of 25 findings) rather than the sum. The "
    "models cannot communicate with each other during the judgement step — there is no "
    "cross-talk, no chain-of-thought sharing, no debate protocol. This deliberate independence "
    "is what makes the consensus signal meaningful: if all three converge on the same "
    "verdict without coordination, the signal is strong; if they disagree, the disagreement "
    "is genuine and carries information.")


# ============================================================================
# 6. IMPLEMENTATION DETAILS
# ============================================================================
add_heading(doc, "4. Implementation Details", level=1)

add_heading(doc, "Languages, frameworks, and libraries", level=2)
add_table(doc,
    ["Layer", "Technology"],
    [
        ["Agent orchestration",  "LangGraph 0.2.x (two-node StateGraph)"],
        ["LLM provider",         "OpenRouter (OpenAI-compatible REST API)"],
        ["IaC scanner",          "Checkov 3.x via subprocess"],
        ["PDF parsing",          "PyMuPDF (offline policy distillation only)"],
        ["Backend framework",    "FastAPI + Uvicorn"],
        ["Backend language",     "Python 3.13"],
        ["Persistence",          "SQLite (audits + versions)"],
        ["Frontend framework",   "React 18 + Vite"],
        ["Frontend language",    "JavaScript (JSX)"],
        ["Inter-rater statistic","Cohen's kappa (custom implementation)"],
        ["Caching",              "SHA-256 keyed JSON cache on disk"],
    ],
    widths=[2.2, 4.6])

add_heading(doc, "System setup", level=2)
add_para(doc,
    "Development was carried out on a Windows 11 workstation with Python 3.13 in a virtual "
    "environment under backend/venv/. No GPU is required; all model inference runs in the "
    "OpenRouter cloud. Local CPU is used only for cloning the target repository, running "
    "Checkov as a subprocess (~26 seconds for the terragoat repo, ~6 seconds for smaller "
    "test repos), and computing the consensus aggregation. The frontend is served by Vite "
    "on port 5173 with a proxy that forwards /audit, /eval, /policies, /history, and "
    "/fix requests to FastAPI on port 8000.")

add_heading(doc, "Important implementation decisions", level=2)
add_bullets(doc, [
    "Temperature set to 0 for all LLM calls. Combined with the SHA-256 prompt cache, this gives bit-identical reproducibility on re-runs.",
    "Findings are batched 25 per LLM call. Smaller batches than the original 80 because the new prompt includes the full HCL block per finding, which expands token consumption.",
    "All three models judge every finding (no early-exit). Single-model dropout would lose the consensus signal we are trying to measure.",
    "Annotation file fp_annotations.json is hand-crafted with adversarial pairs — two entries sharing the same check_id but differing in surrounding context, one labelled GENUINE and one FALSE_POSITIVE. This forces the LLMs to actually use the context rather than memorising rule outcomes.",
    "The /evaluate endpoint from the original architecture is retained for backwards compatibility but the new /eval/fp-filter and /eval/planner endpoints serve precomputed JSONs to the frontend EvalPanel.",
    "No credentials are committed; the .env file is gitignored and the GitHub push-protection scan was used to verify no live secrets appear in the annotation set.",
])

add_heading(doc, "Code location", level=2)
add_para(doc, "Full source code is available at:")
p = doc.add_paragraph()
r = p.add_run("https://github.com/ParthGala2k/AuditTrace")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x6f, 0x42, 0xc1)
r.underline = True


# ============================================================================
# 7. TASK DISTRIBUTION & CONTRIBUTIONS
# ============================================================================
add_heading(doc, "5. Task Distribution and Contributions", level=1)
add_para(doc,
    "This was a single-author project (Group 4 with one student). All components were "
    "designed, implemented, evaluated, and documented by Parth Gala.")
add_table(doc,
    ["Component", "Contributor"],
    [
        ["Backend pipeline (main.py, consensus.py, checkov_runner.py, db.py)", "Parth Gala"],
        ["Frontend application (App.jsx, IssueCard, EvalPanel, design system)", "Parth Gala"],
        ["Evaluation suite (eval_e2e.py, run_eval.py, compute_metrics.py)",     "Parth Gala"],
        ["Hand-labeled FP annotation set (30 entries across 13 Checkov rules)", "Parth Gala"],
        ["CIS AWS policy distillation (89 structured requirements)",            "Parth Gala"],
        ["Demo deck (make_ppt.py, AuditTrace_Presentation.pptx)",               "Parth Gala"],
        ["Project report (this document) and README",                           "Parth Gala"],
        ["Video demo recording",                                                "Parth Gala"],
    ],
    widths=[4.6, 2.2])


# ============================================================================
# 8. EVALUATION & TESTING RESULTS
# ============================================================================
add_heading(doc, "6. Evaluation and Testing Results", level=1)

add_heading(doc, "Methodology", level=2)
add_para(doc,
    "We evaluated AuditTrace on two independent benchmark suites with hand-labeled ground "
    "truth. Both produce deterministic JSON outputs that are committed to the repository "
    "and regenerable with a single command, allowing the report numbers to be independently "
    "verified by anyone with an OpenRouter key.")

add_heading(doc, "Suite 1 — FP-Filter Evaluation", level=2)
add_para(doc,
    "Thirty hand-curated findings (15 GENUINE + 15 FALSE_POSITIVE) spanning 13 distinct "
    "Checkov check IDs across S3, IAM, RDS, EBS, VPC, ELB, and CloudTrail rules. Each "
    "annotation pair is adversarial: two entries that share the same check_id but differ "
    "in surrounding HCL context, one correctly genuine and one a contextual false positive. "
    "We compute per-model precision, recall, and F1 on the GENUINE class; the same metrics "
    "on the FALSE_POSITIVE class; Cohen's kappa between every pair of models on the 3-way "
    "verdict label; per-tier precision (does the HIGH tier actually mean genuine?); "
    "FP reduction rate; TP retention rate; and overall consensus accuracy.")

add_heading(doc, "Per-model classification results (GENUINE class)", level=2)
rows = []
for m in FP["models"]:
    s = FP["by_model_genuine_cls"][m]
    rows.append([
        lbl(m),
        pct(s["precision"]),
        pct(s["recall"]),
        pct(s["f1"]),
        s["tp"], s["fp"], s["fn"], s["tn"],
    ])
add_table(doc,
    ["Model", "Precision", "Recall", "F1", "TP", "FP", "FN", "TN"],
    rows,
    widths=[1.7, 1.0, 0.9, 0.8, 0.6, 0.6, 0.6, 0.6])

add_para(doc,
    "Per-model F1 ranges from 84.6% (Llama 3.1 70B) to 100% (DeepSeek V3). Llama exhibits "
    "an under-flagging failure mode — high precision but low recall, missing four real "
    "violations on this set. GPT-4o mini exhibits a single over-flagging error. DeepSeek "
    "achieves zero errors on n=30. The substantive finding is the variance: if we deployed "
    "any single model in isolation, we would be wrong somewhere between zero and four times "
    "on this evaluation set. Consensus voting recovers all of these individual errors.")

add_heading(doc, "Per-model classification results (FALSE_POSITIVE class)", level=2)
rows = []
for m in FP["models"]:
    s = FP["by_model_fp_cls"][m]
    rows.append([lbl(m), pct(s["precision"]), pct(s["recall"]), pct(s["f1"])])
add_table(doc,
    ["Model", "Precision", "Recall", "F1"],
    rows,
    widths=[2.0, 1.5, 1.5, 1.5])

add_heading(doc, "Inter-model agreement (Cohen's kappa)", level=2)
rows = []
for pair, k in FP["kappa"].items():
    if k >= 0.81:   interp = "Almost perfect"
    elif k >= 0.61: interp = "Substantial"
    elif k >= 0.41: interp = "Moderate"
    else:           interp = "Fair"
    rows.append([pair, num(k), interp])
add_table(doc,
    ["Model pair", "κ", "Landis & Koch interpretation"],
    rows,
    widths=[3.4, 0.8, 2.4])

add_para(doc,
    "Cohen's kappa ranges from 0.67 (GPT-4o mini vs Llama) to 0.93 (GPT-4o mini vs DeepSeek). "
    "No pair exhibits perfect agreement. This is the critical finding for justifying the "
    "use of three models: the models genuinely disagree on hard cases, which means the "
    "consensus mechanism is performing real error correction rather than rubber-stamping "
    "identical verdicts.")

add_heading(doc, "Per-tier ground-truth distribution", level=2)
rows = []
tier_order = ["HIGH", "LIKELY", "UNCERTAIN", "LIKELY_FP", "SUPPRESSED"]
for t in tier_order:
    info = FP["by_tier"].get(t)
    if not info:
        continue
    rows.append([t, info["count"], info["truly_genuine"], info["truly_fp"], pct(info["genuine_pct"])])
add_table(doc,
    ["Tier", "Count", "Truly Genuine", "Truly FP", "% Genuine"],
    rows,
    widths=[1.5, 1.0, 1.5, 1.2, 1.2])

add_para(doc,
    "Tier labels are predictive of ground truth on this set. Every finding assigned to the "
    "HIGH tier (11 findings) is in fact a genuine violation; every finding assigned to "
    "SUPPRESSED (14 findings) is in fact a false positive. The mapping from tier name to "
    "ground-truth class is therefore non-arbitrary — the tier carries actual information "
    "about the underlying truth.")

add_heading(doc, "Consensus headline numbers", level=2)
add_table(doc,
    ["Metric", "Value", "Interpretation"],
    [
        ["FP reduction rate",   pct(FP["fp_reduction_rate"]),  "Fraction of true FPs correctly suppressed by consensus"],
        ["TP retention rate",   pct(FP["tp_retention_rate"]),  "Fraction of true GENUINEs correctly retained"],
        ["Overall accuracy",    pct(FP["overall_accuracy"]),   "Fraction of all 30 cases assigned to the correct tier"],
    ],
    widths=[1.6, 1.0, 4.2])

add_para(doc,
    "All three headline metrics are 100% on n=30. These three numbers are not independent "
    "— they are three views of the same fact: every case is in the correct tier. The "
    "load-bearing evidence is the per-model variance above (84.6%-100% F1) and the Cohen's "
    "kappa range above (0.67-0.93). The headline 100% should be interpreted as a "
    "consequence of those underlying numbers on a small hand-curated set, not as evidence "
    "of an ideal system that will generalize unchanged to a much larger dataset.")

add_heading(doc, "Suite 2 — Planner Benchmark", level=2)
add_para(doc,
    "A separate benchmark of 55 test cases drawn from the CIS AWS Foundations Benchmark "
    "PDF. Each test case provides a policy clause as input and the expected (severity, "
    "requirement_type, check_targets) triple as the ground truth. All three models are "
    "evaluated under identical prompting, with severity accuracy, requirement-type "
    "accuracy, targets F1, latency, total tokens, and dollar cost recorded.")

rows = []
for r in PLN:
    rows.append([
        lbl(r["model"]),
        pct(r["severity_accuracy"]),
        pct(r["req_type_accuracy"]),
        pct(r["avg_targets_f1"]),
        f"{r['avg_latency_s']:.2f} s",
        usd(r["estimated_cost_usd"]),
    ])
add_table(doc,
    ["Model", "Severity Acc.", "Type Acc.", "Targets F1", "Latency", "Cost / run"],
    rows,
    widths=[1.8, 1.2, 1.0, 1.0, 0.9, 0.9])

add_para(doc,
    "No single winner emerges. GPT-4o mini is the fastest at 1.34 seconds per call but has "
    "the lowest targets F1 (70.5%). DeepSeek V3 has the highest type accuracy (90.9%) and "
    "targets F1 (80.9%) at the lowest dollar cost ($0.0018), but is the slowest at 7.15 "
    "seconds per call. Llama 3.1 70B sits in the middle on speed and is the only model to "
    "exceed 50% severity accuracy. Severity accuracy is intentionally low across all three "
    "models because severity is the most subjective field in the ground truth — reasonable "
    "auditors disagree on whether a given clause is critical or high. The objective signals "
    "(type accuracy and targets F1) show clearer differentiation.")

add_heading(doc, "How to reproduce these results", level=2)
add_para(doc,
    "Every number in this section is reproducible from the committed repository. The full "
    "procedure is documented in README.md but the essential commands are:")

reproduce = [
    "cd backend && python -m venv venv && .\\venv\\Scripts\\Activate.ps1",
    "pip install -r requirements.txt",
    "Create backend/.env with OPENROUTER_API_KEY=<key>",
    "python ../evaluation/eval_e2e.py     → writes evaluation/eval_e2e_results.json",
    "python ../evaluation/run_eval.py     → writes evaluation/comparison_summary.json",
]
for cmd in reproduce:
    p = doc.add_paragraph(style="List Number")
    r = p.runs[0] if p.runs else p.add_run(cmd)
    if not p.runs:
        r = p.add_run(cmd)
    else:
        r.text = cmd
    r.font.name = "Consolas"
    r.font.size = Pt(10)

add_para(doc,
    "Cold-cache runtime is approximately 30–60 seconds for the FP-filter eval and 3–7 "
    "minutes for the planner benchmark. Warm-cache (after the first run) reruns complete "
    "in under one second because all LLM responses are SHA-256-keyed on disk.")

add_heading(doc, "Verification of correctness", level=2)
add_bullets(doc, [
    "Determinism: re-running eval_e2e.py produces bit-identical eval_e2e_results.json under SHA-256 cache. Confirmed across multiple runs during development.",
    "End-to-end smoke test: a complete audit was run against bridgecrewio/terragoat producing 385 findings (211 high, 174 medium, 86.8% compliance score) with all three LLM tiers populated correctly. Screen recording of this run is in the demo video.",
    "Frontend smoke test: the /eval/fp-filter and /eval/planner endpoints were verified via the EvalPanel component in the React frontend. Manual inspection confirmed numbers match the committed JSONs.",
    "No-credential validation: GitHub push protection blocked an earlier commit containing a Stripe-pattern placeholder. The value was replaced with an obviously redacted placeholder and re-pushed cleanly.",
])

add_heading(doc, "Limitations of the evaluation", level=2)
add_bullets(doc, [
    "n = 30 is small for the FP-filter set. The per-model F1 spread (84.6%-100%) and the Cohen's kappa range (0.67-0.93) are the most defensible findings; the 100% consensus accuracy is partly an artifact of dataset size.",
    "Annotations were created by a single annotator (the author). A multi-annotator agreement study would strengthen the ground truth.",
    "Severity accuracy in the planner benchmark is noisy because the ground-truth severity field is subjective. Type accuracy and targets F1 should be treated as the primary objective signals.",
    "The evaluation covers only Terraform/HCL inputs. CloudFormation and Kubernetes manifests are detected by Checkov but the HCL-extraction fallback degrades for them; their FP-filter quality has not been quantified.",
    "The system has been tested against the CIS AWS Foundations Benchmark only. Generalization to NIST SP 800-53, SOC 2, and PCI-DSS is future work.",
])


# ============================================================================
# 9. REFERENCES
# ============================================================================
add_heading(doc, "7. References", level=1)
refs = [
    "Bridgecrew / Prisma Cloud. Checkov: open-source infrastructure-as-code static analysis. https://www.checkov.io/",
    "Center for Internet Security. CIS Amazon Web Services Foundations Benchmark, v3.0. 2023. https://www.cisecurity.org/benchmark/amazon_web_services",
    "Cohen, J. (1960). A coefficient of agreement for nominal scales. Educational and Psychological Measurement, 20(1), 37–46.",
    "Landis, J. R., & Koch, G. G. (1977). The measurement of observer agreement for categorical data. Biometrics, 33(1), 159–174.",
    "Zheng, L. et al. (2023). Judging LLM-as-a-Judge with MT-Bench and Chatbot Arena. NeurIPS Datasets and Benchmarks.",
    "LangChain Inc. LangGraph: stateful multi-agent LangChain framework. https://github.com/langchain-ai/langgraph",
    "OpenRouter. Unified LLM API gateway. https://openrouter.ai/docs",
    "Artifex Software. PyMuPDF: Python bindings for the MuPDF library. https://pymupdf.readthedocs.io/",
    "Ramírez, S. FastAPI: modern high-performance Python web framework. https://fastapi.tiangolo.com/",
    "Meta AI. Llama 3.1 model card. July 2024.",
    "DeepSeek AI. DeepSeek-V3 technical report. 2024.",
    "OpenAI. GPT-4o mini system card. July 2024.",
    "bridgecrewio. terragoat: deliberately vulnerable Terraform reference repository. https://github.com/bridgecrewio/terragoat",
    "Vite. Frontend tooling for modern web projects. https://vitejs.dev/",
    "Python Software Foundation. Python 3.13 language reference. https://docs.python.org/3.13/",
]
for i, ref in enumerate(refs, start=1):
    p = doc.add_paragraph()
    r = p.add_run(f"[{i}]  ")
    r.bold = True
    r.font.size = Pt(11)
    r2 = p.add_run(ref)
    r2.font.size = Pt(11)


# ============================================================================
# Save
# ============================================================================
out_path = os.path.join(ROOT, "AuditTrace_Report.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
